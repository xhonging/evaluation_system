import os
import pandas as pd
import numpy as np
from datetime import timedelta
from flask import jsonify
from flask import make_response, send_from_directory

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm
from docx.shared import RGBColor
from docx.oxml.ns import qn

from app.common.function import get_url_data, set_index, set_columns, get_valid_date
from .config import parameters as pt
from .basic.battery import BatteryCluster

import matplotlib.pyplot as plt
from matplotlib.font_manager import FontProperties
from pandas.plotting import register_matplotlib_converters

font = FontProperties(fname=r"c:\windows\fonts\SimSun.ttc", size=14)
register_matplotlib_converters()


def get_reduce_data(start_, end_):
    """
    获取降需数据
    :param start_: 开始时间
    :param end_: 结束时间
    :return: 降需数据response
    """
    title = '[{},{})'.format(start_, end_) + pt.project + '降需数据.xlsx'
    dir_up = os.path.dirname(os.path.abspath(__file__))
    dir_ = os.path.join(dir_up, 'excel')
    directory = os.path.join(dir_, title)

    if os.path.exists(directory):
        return download(dir_, title)
    else:
        start = get_valid_date(start_)
        end = get_valid_date(end_)
        df_data = get_power_data(start, end)
        data_max_power = get_max_power(df_data)
        df_over_power = get_over_time_data(df_data)
        data_dis_end = cal_dis_time(df_data)
        data_part_dis = get_part_dis(df_data)

        with pd.ExcelWriter(directory) as writer:
            data_max_power.to_excel(writer, sheet_name='主变功率最大值')
            df_over_power.to_excel(writer, sheet_name='超出阈值时段')
            data_dis_end.to_excel(writer, sheet_name='放电截止时间点', index=False)
            data_part_dis.to_excel(writer, sheet_name='各时段放电量')

        return download(dir_, title)


def get_reduce_picture(start_, end_):
    title = '[{},{})'.format(start_, end_) + pt.project + '降需数据可视化图.docx'
    dir_up = os.path.dirname(os.path.abspath(__file__))
    dir_ = os.path.join(dir_up, 'docx')
    directory = os.path.join(dir_, title)

    if os.path.exists(directory):
        return download(dir_, title)
    else:
        # 创建 Document 对象，相当于打开一个 word 文档
        document = Document()
        # 设置整个文档的默认字体
        document.styles['Normal'].font.name = u'宋体'
        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

        section = document.sections[0]  # 获取section对象
        section.left_margin = Cm(1)     # 左页边距
        section.right_margin = Cm(1)    # 右页边距

        # 向文档中添加一个标题，标题级别设置为0级
        head = document.add_heading('', level=0).add_run(title)
        head.font.name = u'宋体'
        head._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        head.font.size = Pt(16)
        head.font.bold = True
        head.font.color.rgb = RGBColor(0, 0, 0)

        start = get_valid_date(start_)
        end = get_valid_date(end_)
        df_data = get_power_data(start, end)
        data_max_power = get_max_power(df_data)
        data_dis_end = cal_dis_time(df_data)
        data_part_dis = get_part_dis(df_data)

        # 添加图片
        data_max_power.index = pd.to_datetime(data_max_power.index)
        fig_max = plt_max_power(data_max_power, dir_up)
        document.add_picture(fig_max, width=Inches(8), height=Inches(5))

        fig_end, fig_day = plt_dis_end(data_dis_end, dir_up)
        document.add_picture(fig_end, width=Inches(8), height=Inches(5))
        document.add_picture(fig_day, width=Inches(8), height=Inches(5))

        data_part_dis.index = pd.to_datetime(data_part_dis.index)
        fig_part = plt_dis_part(data_part_dis, dir_up)
        document.add_picture(fig_part, width=Inches(8), height=Inches(5))

        document.save(directory)
        return download(dir_, title)


def plt_max_power(data, dir_):
    df = data.dropna()
    fig = plt.figure(figsize=(12, 8))
    title = '主变功率最大值'
    plt.title(title, fontproperties=font)
    plt.plot(df.index.day, df, label='主变功率最大值', linewidth=2, color='b', marker='o',
             markerfacecolor='red', markersize=6)
    for a, b in zip(df.index.day, df.values):
        plt.text(a, b[0], '%.2f' % b[0], ha='left', va='bottom', fontsize=15)
    plt.xticks(df.index.day)
    plt.xlabel('日期', fontproperties=font)
    plt.ylabel('功率', fontproperties=font)
    plt.legend(prop=font, loc='best')
    fig_path = os.path.join(dir_, 'figure', title+'.jpg')
    plt.savefig(fig_path)

    return fig_path


def plt_dis_part(data, dir_):
    col = data.columns
    dis_8_12 = data[col[0]].dropna()
    dis_12_17 = data[col[1]].dropna()
    dis_17_21 = data[col[2]].dropna()

    fig = plt.figure(figsize=(12, 8))
    bar_width = 0.3
    N = np.arange(len(dis_8_12.index.day))

    plt.bar(N, dis_8_12, bar_width, label='7-12')
    plt.bar(N + bar_width, dis_12_17, bar_width, label='12-17')
    plt.bar(N + 2*bar_width, dis_17_21, bar_width, label='17-21')
    for a, b in zip(N, dis_8_12):
        plt.text(a, b + 0.0005, '%.0f' % b, ha='center', va='bottom', fontsize=15)
    for a, b in zip(N+bar_width, dis_12_17):
        plt.text(a, b + 0.0005, '%.0f' % b, ha='center', va='bottom', fontsize=15)
    for a, b in zip(N + 2*bar_width, dis_17_21):
        plt.text(a, b + 0.0005, '%.0f' % b, ha='center', va='bottom', fontsize=15)

    plt.xticks(N + bar_width, dis_8_12.index.day)
    title = '各时段放电量'
    plt.title(title, fontproperties=font)
    plt.xlabel('日期', fontproperties=font)
    plt.ylabel('放电量', fontproperties=font)
    plt.legend(prop=font, loc='best')
    gif_dir = os.path.join(dir_, 'figure', title + '.jpg')
    plt.savefig(gif_dir)
    # plt.show()
    return gif_dir


def plt_dis_end(data, dir_):
    col = data.columns
    df_time = data[[col[0], col[2], col[4]]].max(axis=1)
    data['time'] = pd.DatetimeIndex(df_time).strftime('%H:%M:%S')
    data['dis'] = data[[col[1], col[3], col[5]]].sum(axis=1)

    df = data[['time', 'dis']]
    df.index = pd.DatetimeIndex(df_time).strftime('%Y-%m-%d')
    df.index = pd.DatetimeIndex(df.index)
    df = df.dropna()

    fig = plt.figure(figsize=(12, 8))
    title = '放电截止时间'
    plt.title(title, fontproperties=font)
    plt.plot(df.index.day, pd.DatetimeIndex(df['time']), label='放电截止时间', linewidth=2, color='b', marker='o',
             markerfacecolor='red', markersize=6)
    for a, b in zip(df.index.day, df['time'].values):
        plt.text(a, b, b, ha='center', va='bottom', fontsize=15)
    plt.xticks(df.index.day)
    plt.xlabel('日期', fontproperties=font)
    plt.ylabel('截止时间', fontproperties=font)
    plt.legend(prop=font, loc='best')
    dir_1 = os.path.join(dir_, 'figure', title+'.jpg')
    plt.savefig(dir_1)
    # plt.show()

    fig2 = plt.figure(figsize=(12, 8))
    title2 = '日放电量'
    plt.title(title2, fontproperties=font)
    bar_width = 0.4
    N = np.arange(len(df.index.day))

    plt.bar(N, df['dis'], bar_width, label='总放电量')
    for a, b in zip(N, df['dis']):
        plt.text(a, b + 0.0005, '%.0f' % b, ha='center', va='bottom', fontsize=15)

    plt.xticks(N, df.index.day)
    plt.xlabel('日期', fontproperties=font)
    plt.ylabel('截止时间', fontproperties=font)
    plt.legend(prop=font, loc='best')
    dir_2 = os.path.join(dir_, 'figure', title2 + '.jpg')
    plt.savefig(dir_2)
    # plt.show()
    return dir_1, dir_2


def get_max_power(df_data):
    """
    计算每日主变功率最大值
    :param df_data: 原始数据
    :return:
    """
    col = df_data.columns
    df_max = pd.DataFrame(columns=['主变功率最大值'], index=df_data.index.strftime('%Y-%m-%d'))
    for d, df in df_data.groupby(df_data.index.strftime('%Y-%m-%d')):
        df_max['主变功率最大值'][d] = df[col[-2]].max()
    df_max = df_max.drop_duplicates()
    return df_max


def get_over_time_data(df_data):
    """
    计算超出阈值时间段
    :param df_data: 原始数据
    :return:
    """
    up_value = 500
    col = df_data.columns
    df_up = df_data.loc[df_data[col[-2]] > up_value]
    df_over_power = pd.DataFrame(columns=['主变功率', '积美负荷功率', 'PCS功率'], index=df_up.index)
    df_over_power['主变功率'] = df_up['主变测控_P']
    df_over_power['积美负荷功率'] = df_up['积美测控_P']
    df_over_power['PCS功率'] = df_up['PCS1_有功功率'] + df_up['PCS2_有功功率'] + df_up['PCS3_有功功率']
    return df_over_power


def cal_dis_time(df_data):
    """
    计算放电截止时间
    :param df_data:
    :return:
    """
    col_dis_max = ['pcs1截止时间', 'pcs1放电量', 'pcs2截止时间', 'pcs2放电量', 'pcs3截止时间', 'pcs3放电量']
    df_dis_max = pd.DataFrame(columns=col_dis_max)
    for d, df in df_data.groupby(df_data.index.strftime('%Y-%m-%d')):
        dis1 = df['PCS1_日放电量']
        dis1_max = df['PCS1_日放电量'][-1]
        dis2 = df['PCS2_日放电量']
        dis2_max = df['PCS2_日放电量'][-1]
        dis3 = df['PCS3_日放电量']
        dis3_max = df['PCS3_日放电量'][-1]

        pcs1 = dis1[dis1 == dis1_max]
        pcs2 = dis2[dis2 == dis2_max]
        pcs3 = dis3[dis3 == dis3_max]
        df_dis_max.loc[len(df_dis_max)] = [pcs1.index[0], pcs1.values[0],
                                           pcs2.index[0], pcs2.values[0],
                                           pcs3.index[0], pcs3.values[0]]

    return df_dis_max


def get_part_dis(df_data):
    """
    分段放电量汇总
    :param df_data:
    :return:
    """
    t1 = ' 07:00:00'
    t2 = ' 12:00:00'
    t3 = ' 17:00:00'
    t4 = ' 21:00:00'
    data_7_12 = get_part_dis_power(df_data, t1, t2)
    data_12_17 = get_part_dis_power(df_data, t2, t3)
    data_17_21 = get_part_dis_power(df_data, t3, t4)

    df_all = pd.DataFrame(columns=['7-12放电量', '12-17放电量', '17-21放电量'])
    df_all['7-12放电量'] = data_7_12['总放电量'].dropna()
    df_all['12-17放电量'] = data_12_17['总放电量'].dropna()
    df_all['17-21放电量'] = data_17_21['总放电量'].dropna()

    df_all.index = df_all.index.strftime('%Y-%m-%d')
    return df_all


def get_part_dis_power(df_data, t1, t2):
    """
    分段放电量
    :param df_data:
    :param t1: 开始时间
    :param t2: 结束时间
    :return:
    """
    col_1 = ['积美负荷功率最大值', '总放电量', 'PCS1_放电量', 'PCS2_放电量', 'PCS3_放电量']
    date_range = [d for d in pd.date_range(pt.first_start, pt.last_end, freq='d', closed='left')]
    df_part_power = pd.DataFrame(columns=col_1, index=date_range)
    for d, df in df_data.groupby(df_data.index.strftime('%Y-%m-%d')):
        s1 = d + t1
        s2 = d + t2
        df1 = df.loc[s1:s2]
        max_power = df1['积美测控_P'].max()
        pcs1_dis = df1['PCS1_日放电量'].max() - df1['PCS1_日放电量'].min()
        pcs2_dis = df1['PCS2_日放电量'].max() - df1['PCS2_日放电量'].min()
        pcs3_dis = df1['PCS3_日放电量'].max() - df1['PCS3_日放电量'].min()
        sum_dis = pcs1_dis + pcs2_dis + pcs3_dis
        df_part_power.loc[d] = [max_power, sum_dis, pcs1_dis, pcs2_dis, pcs3_dis]
    # df_part_power.to_excel(os.path.join(root_dir, t1[1:3] + '~' + t2[1:3] + '点放电量.xlsx'))
    return df_part_power


def get_power_data(start, end):
    ids = [["analoginput", "{}{}_有功功率".format(pt.prefix, 'PCS1')],
           ["analoginput", "{}{}_日充电量".format(pt.prefix, 'PCS1')],
           ["analoginput", "{}{}_日放电量".format(pt.prefix, 'PCS1')],
           ["analoginput", "{}{}_电池组总电流".format(pt.prefix, 'PCS1')],
           ["analoginput", "{}{}_电池总SOC".format(pt.prefix, 'PCS1')],
           ["analoginput", "{}{}_有功功率".format(pt.prefix, 'PCS2')],
           ["analoginput", "{}{}_日充电量".format(pt.prefix, 'PCS2')],
           ["analoginput", "{}{}_日放电量".format(pt.prefix, 'PCS2')],
           ["analoginput", "{}{}_电池组总电流".format(pt.prefix, 'PCS2')],
           ["analoginput", "{}{}_电池总SOC".format(pt.prefix, 'PCS2')],
           ["analoginput", "{}{}_有功功率".format(pt.prefix, 'PCS3')],
           ["analoginput", "{}{}_日充电量".format(pt.prefix, 'PCS3')],
           ["analoginput", "{}{}_日放电量".format(pt.prefix, 'PCS3')],
           ["analoginput", "{}{}_电池组总电流".format(pt.prefix, 'PCS3')],
           ["analoginput", "{}{}_电池总SOC".format(pt.prefix, 'PCS3')],
           ["analoginput", "{}主变测控_P".format(pt.prefix)],
           ["analoginput", "{}积美测控_P".format(pt.prefix)]
           ]

    df = get_url_data(pt.url, ids, pt.func, start, end, pt.interval)
    df.index = set_index(df)
    df.columns = set_columns(df, pt.prefix)

    return df


def download(dir_, title):
    try:
        response = make_response(send_from_directory(dir_, title, as_attachment=True))
        return response
    except Exception as e:
        return jsonify({"code": "异常", "message": "{}".format(e)})
