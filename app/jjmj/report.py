import os
import pandas as pd
from datetime import datetime, timedelta, date
from dateutil.relativedelta import relativedelta
from flask import jsonify
from flask import make_response, send_from_directory

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm
from docx.shared import RGBColor
from docx.oxml.ns import qn

from .config import parameters as pt
from .basic.station import Station
from app.common.function import get_valid_date


def generate_report(start_, end_):
    """
    生成月度报表
    :param start_:开始时间
    :param end_: 结束时间
    :return: 报表
    """
    title = start_[:7] + pt.project + '储能电站运行指标月报表' + '.docx'
    dir_up = os.path.dirname(os.path.abspath(__file__))
    dir_ = os.path.join(dir_up, 'docx')
    directory = os.path.join(dir_, title)

    if os.path.exists(directory):
        return download(dir_, title)
    else:
        # 创建 Document 对象，相当于打开一个 word 文档
        document = Document()
        # 设置整个文档的默认字体
        document.styles['Normal'].font.name = u'宋体'
        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

        # 向文档中添加一个标题，标题级别设置为0级
        head = document.add_heading('', level=0).add_run('电化学储能电站运行指标月报表')
        head.font.name = u'宋体'
        head._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        head.font.size = Pt(16)
        head.font.bold = True
        head.font.color.rgb = RGBColor(0, 0, 0)

        # 向文档中添加一个段落，并将段落引用赋给变量 p
        p = document.add_paragraph('电站名称：')
        p.add_run(pt.project)
        p.add_run(' '*24 + '统计时段：' + ' '*2)
        p.add_run(start_[:4] + '年')
        p.add_run(' '*2 + start_[5:7] + '月')
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        # 添加表格，填入表格内容
        table = document.add_table(rows=30, cols=6, style='Table Grid')
        table.style.font.size = Pt(13)
        table.style.font.name = '宋体'
        table.style.font.color.rgb = RGBColor(0, 0, 0)
        table.style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        table.cell(0, 0).merge(table.cell(0, len(table.columns) - 1))
        table.cell(0, 0).text = '储能电站运行指标统计'

        # 标题栏
        h1 = table.rows[1].cells
        h1[0].text = '指标'
        h1[1].text = '月计划'
        h1[2].text = '月实际'
        h1[3].text = '完成计划/%'
        h1[4].text = '环比/%'
        h1[5].text = '同比/%'

        # 指标项
        # 各时期电站数据获取
        start = get_valid_date(start_)
        end = get_valid_date(end_)
        station = Station(pt.project, start, end)

        lm_start = start - relativedelta(months=1)
        lm_end = start
        lm_station = Station(pt.project, lm_start, lm_end)
        lm_title = lm_start.strftime('%Y-%m-%d')[:7] + pt.project + '储能电站运行指标月报表'
        lm_docx = os.path.join(dir_up, 'docx', lm_title + '.docx')
        if os.path.exists(lm_docx):
            lm_document = Document(lm_docx)
            lm_table = lm_document.tables[0]
        else:
            lm_table = None

        ly_start = start - relativedelta(years=1)
        ly_end = pd.to_datetime(date(end.year - 1, end.month, end.day))
        ly_station = Station(pt.project, ly_start, ly_end)
        ly_title = ly_start.strftime('%Y-%m-%d')[:7] + pt.project + '储能电站运行指标月报表'
        ly_docx = os.path.join(dir_up, 'docx', ly_title + '.docx')

        # 指标
        r0 = table.columns[0].cells
        r0[2].text = '运行时长/天'
        table.cell(2, 2).text = str(station.get_work_days())
        # if lm_table:
        #     lm = float(lm_table.cell(2, 2).text)
        #     bm = float(table.cell(2, 2).text)
        #     table.cell(2, 4).text = str(round((bm - lm) / lm * 110, 2))

        r0[3].text = '下网电量/kWh'
        table.cell(3, 2).text = str(station.get_off_grid_energy())
        if lm_table:
            lm = float(lm_table.cell(3, 2).text)
            bm = float(table.cell(3, 2).text)
            table.cell(3, 4).text = str(round((bm - lm) / lm * 110, 2))

        r0[4].text = '上网电量/kWh'
        table.cell(4, 2).text = str(station.get_on_grid_energy())
        if lm_table:
            lm = float(lm_table.cell(4, 2).text)
            bm = float(table.cell(4, 2).text)
            table.cell(4, 4).text = str(round((bm - lm) / lm * 110, 2))

        r0[5].text = '电站综合效率/%'
        table.cell(5, 2).text = str(station.get_combined_efficiency())
        if lm_table:
            lm = float(lm_table.cell(5, 2).text)
            bm = float(table.cell(5, 2).text)
            table.cell(5, 4).text = str(round((bm - lm) / lm * 110, 2))

        r0[6].text = '直流系统效率/%'
        table.cell(6, 2).text = str(station.get_dc_efficiency())
        if lm_table:
            lm = float(lm_table.cell(6, 2).text)
            bm = float(table.cell(6, 2).text)
            table.cell(6, 4).text = str(round((bm - lm) / lm * 110, 2))

        r0[7].text = 'PCS效率/%'
        table.cell(7, 2).text = str(station.get_ac_efficiency())
        if lm_table:
            lm = float(lm_table.cell(7, 2).text)
            bm = float(table.cell(7, 2).text)
            table.cell(7, 4).text = str(round((bm - lm) / lm * 110, 2))

        r0[8].text = '站用电率/%'
        table.cell(8, 2).text = str(station.get_self_efficiency())
        if lm_table:
            lm = float(lm_table.cell(8, 2).text)
            bm = float(table.cell(8, 2).text)
            table.cell(8, 4).text = str(round((bm - lm) / lm * 110, 2))

        r0[9].text = '经济收益/元'
        table.cell(9, 2).text = str(station.get_earnings())
        if lm_table:
            lm = float(lm_table.cell(9, 2).text)
            bm = float(table.cell(9, 2).text)
            table.cell(9, 4).text = str(round((bm - lm) / lm * 110, 2))

        # r0[10].text = '经济收益/元'
        #
        # r0[11].text = ''
        # table.cell(11, 2).text = ''
        # if lm_table:
        #     lm = float(lm_table.cell(11, 2).text)
        #     bm = float(table.cell(11, 2).text)
        #     table.cell(11, 4).text = str(round((bm - lm) / lm * 110, 2))
        #
        # r0[12].text = '电站储能损耗率/%'
        # table.cell(12, 2).text = str(station.get_loss_efficiency())
        # if lm_table:
        #     lm = float(lm_table.cell(12, 2).text)
        #     bm = float(table.cell(12, 2).text)
        #     table.cell(12, 4).text = str(round((bm - lm) / lm * 110, 2))

        # 设备维护情况
        table.cell(13, 0).merge(table.cell(13, len(table.columns) - 1))
        table.cell(13, 0).text = '设备故障维护统计'

        h1 = table.rows[14].cells
        h1[0].text = '故障时间'
        h1[1].text = '故障描述'
        h1[2].text = '处理情况'
        h1[3].text = '处理人'
        h1[4].text = '恢复时间'
        h1[5].text = '备注'

        table.cell(18, 0).merge(table.cell(18, len(table.columns) - 1))
        table.cell(18, 0).text = '结论'

        # 保存文本
        document.save(directory)
        return download(dir_, title)


def download(dir_, title):
    try:
        response = make_response(send_from_directory(dir_, title, as_attachment=True))
        return response
    except Exception as e:
        return jsonify({"code": "异常", "message": "{}".format(e)})



